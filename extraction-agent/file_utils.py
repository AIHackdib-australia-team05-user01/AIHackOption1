import os
from docx import Document
from openpyxl import load_workbook


def read_file_to_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()

    if ext == ".txt":
        with open(path, "r", encoding="utf-8") as f:
            return f.read()

    elif ext == ".docx":
        doc = Document(path)
        parts = []
        # paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # tables
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text.strip()
                             for cell in row.cells))
        return "\n".join(parts)

    elif ext == ".xlsx":
        wb = load_workbook(path, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"--- Sheet: {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                if any(row):  # skip empty rows
                    row_text = "\t".join(
                        "" if v is None else str(v) for v in row)
                    parts.append(row_text)
        return "\n".join(parts)

    else:
        raise ValueError(f"Unsupported file extension: {ext}")


# Example usage
if __name__ == "__main__":
    text = read_file_to_text("your_file.docx")   # works for .txt, .docx, .xlsx
    print(text[:500])  # preview first 500 chars
